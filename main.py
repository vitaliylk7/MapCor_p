# main.py
"""
MAPCOR_M — Python-версия программы для расчёта ранговых корреляций
( → PySide6 + pandas + numpy + scipy)

Основные функции:
- Загрузка данных (txt/csv/xlsx)
- Выбор признаков через чек-лист (QListWidget с галочками)
- Расчёт Spearman, DIST10, RR (мета-корреляция)
- Отображение исходных данных в QTableView
- Таблица результатов в отдельном окне
- Генерация и просмотр HTML-отчёта
"""

import sys
import os
from pathlib import Path
import datetime
import pandas as pd
import numpy as np
import tkinter as tk
import json
import subprocess

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QGroupBox, QTableView,
    QStatusBar,  QFileDialog, QMessageBox,
    QHeaderView, QDialog, QLabel, QCheckBox, QPushButton, QScrollArea, QGridLayout,
    QDoubleSpinBox, QSpinBox
)
from PySide6.QtCore import Qt, QUrl
from PySide6.QtGui import QStandardItemModel, QStandardItem, QDesktopServices, QFont, QColor

from data import TData
from stat_corr_types import TStatCorr, TExtendedStat
from corr_calculations import calculate_all_correlations

from docx import Document
from docx.shared import Mm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START



# ────────────────────────────────────────────────────────────────
# Вспомогательные функции для цветовой кодировки (должны быть ДО классов!)
# ────────────────────────────────────────────────────────────────
COLOR_SCALE = [
    '#7bb8c2',   #  0 — мягкий бирюзово-голубой (самый лучший)
    '#8cc9d3',   #  1
    '#a0e0e0',   #  2
    '#c2f0eb',   #  3 — очень светлый
    '#e0f7f4',   #  4 — почти белый с оттенком
    '#80cca8',   #  5 — мятный зелёный
    '#a8d97f',   #  6 — светло-салатовый
    '#d9ec5f',   #  7 — лимонно-зелёный

    '#f2e96b',   #  8 — мягкий жёлтый (не кричащий)
    '#f5cf63',   #  9 — жёлто-оранжевый
    '#f7b05b',   # 10 — светлый оранжевый
    '#f28b55',   # 11 — оранжево-красный
    '#e36a52',   # 12 — приглушённый красно-оранжевый
    '#d65a54',   # 13 — бледно-красный (финал, без агрессии)
]

def get_color_index(value, min_val, max_val, median=None):
    """
    Основная функция вычисления индекса цвета (0..13)
    Логика идентична getColorOfMedian в Delphi
    
    Параметры:
        value   — текущее значение
        min_val — минимум диапазона
        max_val — максимум диапазона
        median  — медиана (если None — линейное деление)
    
    Возвращает: индекс 0..13
    """
    if np.isnan(value):
        return 7  # середина, серый
    
    # Защита от выхода за границы
    value = max(min_val, min(max_val, value))
    
    if median is not None:
        # Деление относительно медианы 
        if value >= median:
            # Выше медианы → индексы 7..13 (7 интервалов)
            portion = (value - median) / (max_val - median) if max_val > median else 0
            ind = round(portion * 7) + 6
            ind = min(13, max(7, ind))
        else:
            # Ниже медианы → индексы 0..6
            portion = (value - min_val) / (median - min_val) if median > min_val else 0
            ind = round(portion * 7)
            ind = min(6, max(0, ind))
    else:
        # Линейное деление всего диапазона на 14 частей
        portion = (value - min_val) / (max_val - min_val) if max_val > min_val else 0
        ind = round(portion * 13)
        ind = min(13, max(0, ind))
    
    return ind


def get_color_for_r(value, median=None):
    """Цвет для Spearman R (диапазон -1..1)"""
    return COLOR_SCALE[get_color_index(value, -1.0, 1.0, median)]


def get_color_for_rr(value, median=None):
    """Цвет для RR (мета-корреляция, тоже -1..1)"""
    return COLOR_SCALE[get_color_index(value, -1.0, 1.0, median)]


def get_color_for_dist10(value, median=None):
    """Цвет для DIST10 (0..100)"""
    return COLOR_SCALE[get_color_index(value, 0.0, 100.0, median)]


# ────────────────────────────────────────────────────────────────
# Дальше идут классы и остальной код
# ────────────────────────────────────────────────────────────────

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("MAPCOR — P версия")
        self.resize(1280, 800)
        self.closeEvent = self._on_close

        self.data = TData()
        self.stat_corr = TStatCorr()
        self.associations = None
        # Главный горизонтальный layout (левая + правая часть)
        main_layout = QHBoxLayout()
        central = QWidget()
        central.setLayout(main_layout)
        self.setCentralWidget(central)

        # Левая панель — выбор характеристик
        left_group = QGroupBox("Выбор характеристик")
        left_layout = QVBoxLayout()          # ✅ Без родителя
        left_group.setLayout(left_layout)    # ✅ Явная привязка

        # ─── Кнопки управления выбором ──────────────────────────────────────
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)

        btn_all = QPushButton("Выбрать все")
        btn_none = QPushButton("Снять все")
        btn_inv = QPushButton("Инвертировать")

        btn_all.clicked.connect(lambda: self._set_all_checked(True))
        btn_none.clicked.connect(lambda: self._set_all_checked(False))
        btn_inv.clicked.connect(self._invert_checks)

        btn_layout.addWidget(btn_all)
        btn_layout.addWidget(btn_none)
        btn_layout.addWidget(btn_inv)
        btn_layout.addStretch()

        left_layout.addLayout(btn_layout)

        # ─── Прокручиваемая область с сеткой чекбоксов ──────────────────────
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        self.check_container = QWidget()
        self.grid_layout = QGridLayout(self.check_container)
        self.grid_layout.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        self.grid_layout.setContentsMargins(12, 8, 12, 12)
        self.grid_layout.setSpacing(8)

        self.scroll_area.setWidget(self.check_container)
        left_layout.addWidget(self.scroll_area)

        # Добавляем в основной layout с желаемой шириной
        main_layout.addWidget(left_group, 5)   # ← 5 частей — левая панель шире

        # ─── Правая часть (таблица + настройки под ней) ──────────────────
        right_column = QVBoxLayout()  # вертикальный контейнер для правой стороны

        # Правая панель — исходные данные
        right_group = QGroupBox("Исходные данные")
        right_layout = QVBoxLayout()
        right_group.setLayout(right_layout)

        self.table_view = QTableView()
        self.table_view.setAlternatingRowColors(True)
        self.table_view.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
        self.table_view.verticalHeader().setVisible(False)
        right_layout.addWidget(self.table_view)

        right_column.addWidget(right_group, stretch=1)  # таблица растягивается

        # ─── Секция настроек — под таблицей ──────────────────────────────
        settings_group = QGroupBox("Настройки расчёта")
        settings_layout = QVBoxLayout()
        settings_group.setLayout(settings_layout)

       # ─── Настройки ассоциаций ──────────────────────────────────────────
        assoc_label = QLabel("Настройки формирования ассоциаций")
        assoc_label.setFont(QFont("Segoe UI", 10, QFont.Weight.Bold))
        settings_layout.addWidget(assoc_label)
        settings_layout.addSpacing(8)

        # Порог корреляции с корнем
        h1 = QHBoxLayout()
        lbl1 = QLabel("Порог корреляции с корнем кластера:")
        self.threshold_root_spin = QDoubleSpinBox()
        self.threshold_root_spin.setRange(0.0, 1.0)
        self.threshold_root_spin.setSingleStep(0.05)
        self.threshold_root_spin.setDecimals(2)
        self.threshold_root_spin.setValue(0.80)
        h1.addWidget(lbl1)
        h1.addWidget(self.threshold_root_spin)
        settings_layout.addLayout(h1)

        # Порог средней корреляции внутри кластера
        h2 = QHBoxLayout()
        lbl2 = QLabel("Порог средней корреляции с кластером:")
        self.threshold_avg_spin = QDoubleSpinBox()
        self.threshold_avg_spin.setRange(0.0, 1.0)
        self.threshold_avg_spin.setSingleStep(0.05)
        self.threshold_avg_spin.setDecimals(2)
        self.threshold_avg_spin.setValue(0.30)
        h2.addWidget(lbl2)
        h2.addWidget(self.threshold_avg_spin)
        settings_layout.addLayout(h2)

        # Максимальное количество итераций
        h3 = QHBoxLayout()
        lbl3 = QLabel("Макс. итераций перераспределения:")
        self.max_iters_spin = QSpinBox()
        self.max_iters_spin.setRange(1, 10000)
        self.max_iters_spin.setValue(20)
        h3.addWidget(lbl3)
        h3.addWidget(self.max_iters_spin)
        settings_layout.addLayout(h3)

        # Порог сходимости
        h4 = QHBoxLayout()
        lbl4 = QLabel("Порог сходимости (доля перемещённых):")
        self.convergence_epsilon_spin = QDoubleSpinBox()
        self.convergence_epsilon_spin.setRange(0.0, 1.0)
        self.convergence_epsilon_spin.setSingleStep(0.05)
        self.convergence_epsilon_spin.setDecimals(2)
        self.convergence_epsilon_spin.setValue(0.00)
        h4.addWidget(lbl4)
        h4.addWidget(self.convergence_epsilon_spin)
        settings_layout.addLayout(h4)
       

        # После btn_stats или после чекбокса в settings_layout
        
        right_column.addWidget(settings_group, stretch=0)  # настройки не растягиваются

        # Добавляем правую колонку в главный layout
        main_layout.addLayout(right_column, stretch=4)  # правая часть уже левой


        # Статусбар
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.statusBar.showMessage("Готов к открытию файла...")

        # Создаём меню
        self._create_menu()

        # Загружаем настройки
        self.working_directory = self._load_settings_from_file()
        
        # ─── Подключение сигналов сохранения ───────────────────────────────
        self.threshold_root_spin.valueChanged.connect(self._save_settings)
        self.threshold_avg_spin.valueChanged.connect(self._save_settings)
        self.max_iters_spin.valueChanged.connect(self._save_settings)
        self.convergence_epsilon_spin.valueChanged.connect(self._save_settings)


    def _save_settings(self):
        """Сохраняет все настройки в settings.json"""
        settings = {
            "threshold_root": round(self.threshold_root_spin.value(), 2),
            "threshold_avg": round(self.threshold_avg_spin.value(), 2),
            "max_iters": self.max_iters_spin.value(),
            "convergence_epsilon": round(self.convergence_epsilon_spin.value(), 2)
        }
        try:
            with open("settings.json", "w", encoding="utf-8") as f:
                json.dump(settings, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print("Ошибка сохранения настроек:", e)
            # можно QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить настройки:\n{str(e)}")

    def _get_initial_dir(self):
        """Возвращает начальную директорию для файловых диалогов"""
        # Проверяем, существует ли рабочая директория
        if os.path.isdir(self.working_directory):
            return self.working_directory
        else:
            # Если нет, используем текущую директорию
            return "."


    def _load_settings_from_file(self):
        """Загружает настройки из settings.json и возвращает рабочую директорию"""
        try:
            with open("settings.json", "r", encoding="utf-8") as f:
                settings = json.load(f)
                
                # Загружаем значения параметров
                self.threshold_root_spin.setValue(settings.get("threshold_root", 0.80))
                self.threshold_avg_spin.setValue(settings.get("threshold_avg", 0.30))
                self.max_iters_spin.setValue(settings.get("max_iters", 20))
                self.convergence_epsilon_spin.setValue(settings.get("convergence_epsilon", 0.00))
                
                # Возвращаем рабочую директорию (по умолчанию текущая директория)
                return settings.get("working_directory", ".")
        except (FileNotFoundError, json.JSONDecodeError, KeyError):
            # Возвращаем текущую директорию по умолчанию
            return "."

    def _save_settings(self):
        """Сохраняет все настройки в settings.json"""
        settings = {
            "threshold_root": round(self.threshold_root_spin.value(), 2),
            "threshold_avg": round(self.threshold_avg_spin.value(), 2),
            "max_iters": self.max_iters_spin.value(),
            "convergence_epsilon": round(self.convergence_epsilon_spin.value(), 2),
            "working_directory": self.working_directory
        }
        try:
            with open("settings.json", "w", encoding="utf-8") as f:
                json.dump(settings, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print("Ошибка сохранения настроек:", e)
            # можно QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить настройки:\n{str(e)}")

    # Заглушки для новых методов (реализуем позже)
    def form_associations(self):
        if self.associations is not None:
            reply = QMessageBox.question(
                self, "Пересчитать?",
                "Ассоциации уже сформированы. Пересчитать заново?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                return

        if self.stat_corr.count() == 0:
            QMessageBox.warning(self, "Нет данных", "Сначала рассчитайте корреляции.")
            return

        params = {
            "threshold_root": self.threshold_root_spin.value(),
            "threshold_avg": self.threshold_avg_spin.value(),
            "max_iters": self.max_iters_spin.value(),
            "convergence_epsilon": self.convergence_epsilon_spin.value()
        }

        from associations import build_associations   # предполагаем, что файл associations.py в той же папке

        self.associations = build_associations(self.stat_corr, params)

        if not self.associations:
            QMessageBox.information(self, "Результат", "Ассоциации не сформированы (слишком слабые связи).")
        else:
            QMessageBox.information(self, "Готово", f"Сформировано {len(self.associations)} ассоциаций/одиночных групп.")



    def generate_assoc_report_docx(self):
        """
        Создаёт и сохраняет отчёт по ассоциациям в формате .docx
        """
        if self.associations is None or not self.associations:
            QMessageBox.information(
                self,
                "Нет данных",
                "Сначала сформируйте ассоциации (меню Ассоциации → Сформировать ассоциации)."
            )
            return

        # Диалог сохранения
        default_name = "Ассоциации_" + datetime.datetime.now().strftime("%Y-%m-%d_%H-%M") + ".docx"
        import os
        tempfname = os.path.join(self._get_initial_dir(), default_name)
        fname, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить отчёт по ассоциациям",     
            tempfname,
            "Word документы (*.docx);;Все файлы (*.*)"
        )
        if not fname:
            return
        if not fname.lower().endswith('.docx'):
            fname += '.docx'

        # Обновляем рабочую директорию на директорию сохраненного файла
        self.working_directory = os.path.dirname(fname)

        # Создаём документ
        doc = Document()

        # Стили
        style_normal = doc.styles['Normal']
        font = style_normal.font
        font.name = 'Arial'
        font.size = Pt(11)

        style_heading1 = doc.styles['Heading 1']
        style_heading1.font.name = 'Arial'
        style_heading1.font.size = Pt(16)
        style_heading1.font.bold = True

        style_heading2 = doc.styles['Heading 2']
        style_heading2.font.name = 'Arial'
        style_heading2.font.size = Pt(14)
        style_heading2.font.bold = True

        # Заголовок документа
        title = doc.add_paragraph("Отчёт по ассоциациям признаков")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Heading 1'

        # Подзаголовок с именем файла и датой
        p = doc.add_paragraph()
        p.add_run(f"Файл данных: ").bold = True
        p.add_run(f"{Path(self.data.filename).name if self.data.filename else 'не загружен'}\n")
        p.add_run(f"Дата формирования: ").bold = True
        p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        doc.add_paragraph()  # пустая строка

        # Параметры расчёта
        doc.add_heading("Параметры формирования ассоциаций", level=2)
        params_table = doc.add_table(rows=5, cols=2)
        params_table.style = 'Table Grid'
        params_table.columns[0].width = Inches(2.8)
        params_table.columns[1].width = Inches(3.5)

        rows_data = [
            ("Порог корреляции с корнем", f"{self.threshold_root_spin.value():.2f}"),
            ("Порог средней корреляции с кластером", f"{self.threshold_avg_spin.value():.2f}"),
            ("Максимальное количество итераций", str(self.max_iters_spin.value())),
            ("Порог сходимости (доля перемещённых)", f"{self.convergence_epsilon_spin.value():.2f}"),
            ("Количество ассоциаций / одиночных", f"{len([c for c in self.associations if len(c['features']) > 1])} / {len([c for c in self.associations if len(c['features']) == 1])}")
        ]

        for i, (label, value) in enumerate(rows_data):
            params_table.rows[i].cells[0].text = label
            params_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # ─── Основная часть: ассоциации ─────────────────────────────────────
        doc.add_heading("Ассоциации признаков", level=2)

        assoc_num = 1
        singles = []

        for cluster in self.associations:
            size = len(cluster['features'])
            if size == 1:
                singles.append(cluster)
                continue

            # Заголовок ассоциации
            doc.add_heading(f"Ассоциация {assoc_num}  (размер: {size})", level=3)

            p = doc.add_paragraph()
            p.add_run("Корень кластера: ").bold = True
            p.add_run(self.stat_corr.get_column_name(cluster['root']))

            p = doc.add_paragraph()
            p.add_run("Средняя корреляция внутри ассоциации: ").bold = True
            p.add_run(f"{cluster['internal_avg_r']:.3f}")

            p = doc.add_paragraph()
            p.add_run("Средняя корреляция с внешними признаками: ").bold = True
            p.add_run(f"{cluster['external_avg_r']:.3f}")

            # Список признаков
            doc.add_paragraph("Признаки в ассоциации:", style='List Bullet')
            for idx in cluster['features']:
                name = self.stat_corr.get_column_name(idx)
                doc.add_paragraph(f"• {name}", style='List Bullet')

            doc.add_paragraph()  # разделитель
            assoc_num += 1

        # ─── Одиночные признаки ─────────────────────────────────────────────
        if singles:
            doc.add_heading("Одиночные признаки (не вошедшие в ассоциации)", level=2)

            for cluster in singles:
                name = self.stat_corr.get_column_name(cluster['features'][0])
                p = doc.add_paragraph(f"• {name}")
                p.paragraph_format.left_indent = Inches(0.3)

        # Сохраняем
        try:
            doc.save(fname)
            QMessageBox.information(
                self,
                "Отчёт сохранён",
                f"Отчёт успешно сохранён в:\n{fname}"
            )

            # Опционально: открыть файл
            if QMessageBox.question(
                self, "Открыть файл?",
                "Открыть созданный документ Word?",
                QMessageBox.Yes | QMessageBox.No
            ) == QMessageBox.Yes:
                import os
                os.startfile(fname)  # Windows
                # Для других ОС можно использовать QDesktopServices.openUrl(QUrl.fromLocalFile(fname))

        except Exception as e:
            QMessageBox.critical(
                self,
                "Ошибка сохранения",
                f"Не удалось сохранить отчёт:\n{str(e)}"
            )

    def _on_close(self, event):
        # Сохраняем перед закрытием
        self._save_settings()
        event.accept()

    def get_selected_columns(self) -> list[int]:
        """Возвращает индексы выбранных (и активных) признаков"""
        selected = []
        for i, cb in enumerate(self.check_boxes):
            if cb.isChecked() and cb.isEnabled():
                selected.append(i)
        return selected

    def _set_all_checked(self, checked: bool):
        """Устанавливает состояние всем чекбоксам"""
        for cb in self.check_boxes:
            if cb.isEnabled():  # не трогаем отключённые
                cb.setChecked(checked)


    def _invert_checks(self):
        """Инвертирует состояние всех доступных чекбоксов"""
        for cb in self.check_boxes:
            if cb.isEnabled():
                cb.setChecked(not cb.isChecked())

    def act_save_statistics_ext_to_csv(self):
        """Сохраняет статистику характеристик в CSV"""
        if not hasattr(self, 'data') or self.data.df is None or self.data.df.empty:
            QMessageBox.warning(self, "Нет данных", "Сначала загрузите файл данных.")
            return
        
        success = self.data.save_statistics_to_csv(working_directory=self.working_directory)
        if success:
            QMessageBox.information(self, "Успех", "Статистика характеристик сохранена в CSV-файл.")
        else:
            QMessageBox.critical(self, "Ошибка", "Не удалось сохранить файл статистики.")

    def fill_features_list(self):
        """Заполняет сетку чекбоксов в 3 столбца + подсветка отключённых"""
        
        # 1. Очищаем старую сетку
        while self.grid_layout.count():
            item = self.grid_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.check_boxes = []  # список всех QCheckBox

        features = self.data.get_column_names()
        if not features:
            lbl = QLabel("Нет загруженных признаков")
            self.grid_layout.addWidget(lbl, 0, 0, 1, 3, Qt.AlignCenter)
            return

        COLUMNS = 3  # фиксированное количество столбцов

        for i, col_name in enumerate(features):
            cb = QCheckBox(col_name)
            cb.setChecked(True)

            # ─── Подсветка и отключение проблемных признаков ───────────────
            is_invalid = False
            if hasattr(self.data, 'invalid_columns') and i in self.data.invalid_columns:
                is_invalid = True

            if is_invalid:
                cb.setChecked(False)
                cb.setEnabled(False)
                cb.setStyleSheet("""
                    color: #888888;
                    font-style: italic;
                """)

            row = i // COLUMNS
            col = i % COLUMNS

            self.grid_layout.addWidget(cb, row, col, Qt.AlignLeft | Qt.AlignVCenter)
            self.check_boxes.append(cb)

        # Растягиваем столбцы равномерно
        for c in range(COLUMNS):
            self.grid_layout.setColumnStretch(c, 1)

        # Добавляем пространство внизу
        self.grid_layout.setRowStretch(self.grid_layout.rowCount(), 1)

        # Сбрасываем скроллбар вверх
        self.scroll_area.verticalScrollBar().setValue(0)

    def get_selected_columns(self) -> list[int]:
        """Возвращает индексы выбранных и активных признаков"""
        return [i for i, cb in enumerate(self.check_boxes) if cb.isChecked() and cb.isEnabled()]


    def get_all_columns_count(self) -> int:
        """Количество всех признаков (включая отключённые)"""
        return len(self.check_boxes)

    def _create_menu(self):
        menubar = self.menuBar()

        file_menu = menubar.addMenu("Файл")
        file_menu.addAction("Открыть данные...", self.act_open, shortcut="Ctrl+O")
        file_menu.addAction("Отчёт статистики (Word)", self.act_save_stats_to_word)
        #file_menu.addAction("Сохранить статистику v2 *.csv", self.act_save_statistics_ext_to_csv)
        
        file_menu.addSeparator()
        file_menu.addAction("Выход", self.close, shortcut="Alt+F4")

        correlation_menu = menubar.addMenu("Корреляции")
        correlation_menu.addAction("Вычислить корреляции", self.act_run, shortcut="F9")          
        correlation_menu.addAction("Отчёт корреляций (HTML)", self.act_view_report_ext)
        correlation_menu.addAction("Отчёт корреляций матричный(HTML)", self.act_view_report_old)
        correlation_menu.addAction("Сохранить корреляции в файл...", self.act_save_result)
        
        assoc_menu = menubar.addMenu("Ассоциации")
        assoc_menu.addAction("Сформировать ассоциации", self.form_associations, shortcut="F10")
        assoc_menu.addAction("Отчет ассоциаций (Word)", self.generate_assoc_report_docx)
       
    def act_open(self):
        fname, _ = QFileDialog.getOpenFileName(
            self, "Открыть файл данных",
            self._get_initial_dir(),
            "Данные (*.elnm *.txt *.dat);;Все файлы (*.*)"
        )
        if not fname:
            return

        # Обновляем рабочую директорию на директорию выбранного файла
        self.working_directory = os.path.dirname(fname)

        self.stat_corr.clear()
        self.table_view.setModel(None)

        self.statusBar.showMessage("Загрузка файла... Подождите...")

        success = self.data.load_file(fname)

        if success:
            # Заполняем таблицу данных
            model = QStandardItemModel()
            model.setHorizontalHeaderLabels(self.data.get_column_names())

            for row_idx in range(self.data.get_count_record()):
                for col_idx, col_name in enumerate(self.data.get_column_names()):
                    val = self.data.df.iloc[row_idx, col_idx]
                    if not pd.isna(val):
                        try:
                            # Проверяем, является ли значение числовым
                            float_val = float(val)
                            item = QStandardItem(f"{float_val:.4g}")
                        except (ValueError, TypeError):
                            # Если не число, используем строковое представление
                            item = QStandardItem(str(val))
                    else:
                        item = QStandardItem("—")
                    item.setEditable(False)
                    model.setItem(row_idx, col_idx, item)

            self.table_view.setModel(model)

            # Заполняем чекбоксы в левой панели
            self.fill_features_list()

            self.statusBar.showMessage(
                f"Файл загружен: {Path(fname).name}   •   строк: {self.data.get_count_record()}   •   признаков: {self.data.get_count_column()}"
            )
        else:
            QMessageBox.warning(self, "Ошибка загрузки",
                                "Не удалось загрузить файл.\nПодробности в data_load.log")
            self.statusBar.showMessage("Ошибка загрузки файла")

    def act_run(self):
        """
        Запуск расчёта корреляций по выбранным признакам.
        """
        # Получаем глобальные индексы выбранных и активных столбцов
        selected_global = self.get_selected_columns()

        if len(selected_global) < 2:
            QMessageBox.warning(
                self,
                "Предупреждение",
                f"Для расчёта корреляций необходимо выбрать минимум 2 признака.\n"
                f"Сейчас выбрано: {len(selected_global)}"
            )
            return

        # Очищаем предыдущие результаты корреляций и ассоциаций
        self.stat_corr.clear()
        self.associations = None  # сбрасываем предыдущие ассоциации, т.к. корреляции пересчитываются

        # Имена только выбранных столбцов (в том порядке, в котором они выбраны)
        selected_names = [self.data.get_column_name(idx) for idx in selected_global]

        # Инициализируем TStatCorr ТОЛЬКО выбранными признаками
        self.stat_corr.initialize(selected_names)

        # Переносим информацию о невалидных столбцах (локальные индексы)
        self.stat_corr.invalid_columns = [
            local_idx for local_idx, global_idx in enumerate(selected_global)
            if global_idx in self.data.invalid_columns
        ]

        # Создаём все пары из выбранных столбцов (используем ЛОКАЛЬНЫЕ индексы 0..n-1)
        num_selected = len(selected_global)
        for i in range(num_selected - 1):
            for j in range(i + 1, num_selected):
                self.stat_corr.add_or_get_pair(i, j)  # ← ЛОКАЛЬНЫЕ i и j!

        # Сообщение о начале расчёта
        self.statusBar.showMessage(f"Расчёт корреляций в режиме …")

        # Адаптер для доступа к данным: локальный индекс → значение из исходных данных
        def get_data_adapter(local_col: int, row: int) -> float:
            global_col = selected_global[local_col]
            return self.data.get_data(global_col, row)

        try:
            calculate_all_correlations(
                stat_corr=self.stat_corr,
                get_data=get_data_adapter,
                num_records=self.data.get_count_record(),
                percent10=10
            )

            pair_count = self.stat_corr.count()
            feature_count = len(self.stat_corr.column_names)

            # Успешное завершение
            QMessageBox.information(
                self,
                "Расчёт завершён",
                f"Успешно рассчитано корреляций для {feature_count} признаков.\n"
                f"Количество пар: {pair_count}"
            )

            self.statusBar.showMessage(
                f"Готово. Рассчитано {pair_count} пар для {feature_count} признаков."
            )

        except Exception as e:
            import traceback
            traceback.print_exc()  # для отладки в консоль

            QMessageBox.critical(
                self,
                "Ошибка при расчёте корреляций",
                f"Произошла ошибка во время вычисления корреляций:\n\n{str(e)}\n\n"
                f"Проверьте данные и настройки. Если проблема повторяется — сообщите разработчику."
            )

            self.statusBar.showMessage("Ошибка при расчёте корреляций")

    def act_view_report_old(self):
        """Вариант отчета из старой версии"""
        if self.stat_corr.count() == 0:
            QMessageBox.information(self, "Нет результатов", "Сначала выполните расчёт.")
            return

        html_content = self._generate_old_report()

        report_path = Path(self.working_directory) / "mapcor_report.html"
        report_path.write_text(html_content, encoding="utf-8")

        QDesktopServices.openUrl(QUrl.fromLocalFile(str(report_path.absolute())))

    def act_view_report_ext(self):
        if self.stat_corr.count() == 0:
            QMessageBox.information(self, "Нет результатов", "Сначала выполните расчёт (F9).")
            return

        html_content = self._generate_extended_report()

        report_path = Path(self.working_directory) / "mapcor_extended_report.html"
        report_path.write_text(html_content, encoding="utf-8")

        QDesktopServices.openUrl(QUrl.fromLocalFile(str(report_path.absolute())))

        self.statusBar.showMessage("Расширенный отчёт открыт в браузере")

    def act_save_stats_to_word(self):
        """
        Экспорт статистики в Word A4.
        - Шрифт Times New Roman 14 pt
        - Первый столбец шире
        - Убираем лишние нули после запятой
        - Полностью нулевой столбец → жирный курсив
        - Полностью нулевая строка → жирный курсив (кроме Признака)
        """
        stats_df = self.data.get_full_statistics()
        if stats_df is None or stats_df.empty:
            QMessageBox.information(self, "Нет данных", "Нет статистики для экспорта.")
            return

        # Фильтрация по выбранным признакам
        selected_cols = self.get_selected_columns()
        if selected_cols:
            valid_names = [self.data.get_column_name(i) for i in selected_cols
                           if 0 <= i < len(self.data.df.columns)]
            stats_df = stats_df.loc[stats_df.index.isin(valid_names)]

        if stats_df.empty:
            QMessageBox.information(self, "Нет данных", "Нет выбранных признаков.")
            return

        # Оставляем только нужные столбцы
        desired_columns = [
            'repeating_min_percent', 'min', 'max', 'mean', 'median',
            'std', 'CV_percent', 'J'
        ]
        existing_desired = [c for c in desired_columns if c in stats_df.columns]
        stats_df = stats_df[existing_desired].copy()
        import os
        tempfname = os.path.join(self._get_initial_dir(), "statistics.docx")
        # Диалог сохранения
        fname, _ = QFileDialog.getSaveFileName(
            self, "Сохранить в Word",
            tempfname,
            "Word (*.docx)"
        )
        if not fname:
            return
        if not fname.lower().endswith('.docx'):
            fname += '.docx'

        # Обновляем рабочую директорию на директорию сохраненного файла
        self.working_directory = os.path.dirname(fname)

        try:
            doc = Document()

            # Настройка страницы A4
            section = doc.sections[0]
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.left_margin = Mm(20)
            section.right_margin = Mm(20)
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)

            # Заголовок
            title = doc.add_paragraph("Статистический отчёт по признакам")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True

            # Подзаголовок
            subtitle = doc.add_paragraph(
                f"Файл: {Path(self.data.filename).name}   |   "
                f"Записей: {self.data.get_count_record():,}   |   "
                f"Признаков: {len(stats_df)}"
            )
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.name = 'Times New Roman'
            subtitle.runs[0].font.size = Pt(12)
            subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)

            doc.add_paragraph()

            # Таблица
            table = doc.add_table(rows=1, cols=len(stats_df.columns) + 1)
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False

            # Заголовки
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Хар-ка'

            header_names = {               
                'repeating_min_percent': 'BLS, %',
                'min': 'Min',
                'max': 'Max',
                'mean': 'Mean',
                'median': 'Med',
                'std': 'Std',
                'CV_percent': 'V, %',
                'J': 'J'
            }

            for i, col in enumerate(stats_df.columns, 1):
                hdr_cells[i].text = header_names.get(col, col)
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                        # Данные + форматирование
            for row_idx, (feature, row) in enumerate(stats_df.iterrows()):
                row_cells = table.add_row().cells
                row_cells[0].text = feature
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                for col_idx, col_name in enumerate(stats_df.columns, 1):
                    val = row[col_name]
                    if pd.isna(val):
                        val_str = "—"
                    elif col_name in ['min', 'median', 'max', 'mean', 'std']:
                        val_str = f"{val:g}"
                    elif col_name in ['repeating_min_percent', 'CV_percent']:
                        val_str = f"{val:g}"
                    elif col_name == 'J':
                        val_str = f"{val:g}"
                    else:
                        val_str = str(val)

                    cell = row_cells[col_idx]
                    cell.text = val_str
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    for run in cell.paragraphs[0].runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

            # Применяем жирный курсив ко всему столбцу «Признак» (индекс 0)
            for row in table.rows:
                cell = row.cells[0]
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.italic = True

            # Применяем жирный курсив ко всей строке заголовков (индекс 0)
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.italic = True

            # Ширина столбцов
            table.columns[0].width = Mm(48)  # шире для Признака
            for i in range(1, len(table.columns)):
                table.columns[i].width = Mm(16)

            # Повтор шапки
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            trPr = table.rows[0]._tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), 'true')
            trPr.append(tblHeader)

            # Легенда
            doc.add_paragraph()
            legend_title = doc.add_paragraph("Расшифровка показателей")
            legend_title.runs[0].font.name = 'Times New Roman'
            legend_title.runs[0].font.size = Pt(14)
            legend_title.runs[0].bold = True

            legend_items = [
                "BLS, % — доля значений ниже чувствительности",
                "Min — минимальное значение",
                "Max — максимальное значение",
                "Mean — матожидание",
                "Med — медиана",
                "Std — стандартное отклонение",
                "V, % — коэффициент вариации (станд. отклон. / |mean| × 100)",
                "J  — нормированная энтропия Шеннона (6 интервалов)\n"
                "   · J ≈ 0 — почти все значения в одном интервале\n"
                "   · J ≈ 1 — равномерное распределение\n"
                "   · Порог однородности: J ≥ 0.65"
            ]

            for item in legend_items:
                # Разбиваем элемент на строки по символу новой строки
                lines = item.split('\n')
                first_line = lines[0].rstrip()  # убираем лишние пробелы справа
                
                # Обрабатываем первую строку с разделителем " — "
                if " — " in first_line:
                    short_name, description = first_line.split(" — ", 1)
                    
                    # Создаём абзац со стилем маркированного списка
                    p = doc.add_paragraph(style='List Bullet')
                    
                    # Часть 1: краткое название — жирный курсив
                    run1 = p.add_run(short_name)
                    run1.bold = True
                    run1.italic = True
                    run1.font.name = 'Times New Roman'
                    run1.font.size = Pt(12)
                    run1.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # для кириллицы
                    
                    # Часть 2: разделитель " — " — обычный курсив
                    run2 = p.add_run(" — ")
                    run2.italic = True
                    run2.font.name = 'Times New Roman'
                    run2.font.size = Pt(12)
                    run2.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    
                    # Часть 3: описание — обычный курсив
                    run3 = p.add_run(description)
                    run3.italic = True
                    run3.font.name = 'Times New Roman'
                    run3.font.size = Pt(12)
                    run3.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                else:
                    # Если нет разделителя — весь текст курсивом
                    p = doc.add_paragraph(first_line, style='List Bullet')
                    for run in p.runs:
                        run.italic = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Обрабатываем дополнительные строки (подпункты) с отступом
                for extra_line in lines[1:]:
                    extra_line = extra_line.strip()
                    if extra_line:  # пропускаем пустые строки
                        p_extra = doc.add_paragraph(extra_line)
                        p_extra.paragraph_format.left_indent = Pt(20)  # отступ слева
                        
                        # Весь текст подпункта — курсивом
                        for run in p_extra.runs:
                            run.italic = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            doc.save(fname)
            QMessageBox.information(self, "Сохранено", f"Отчёт сохранён:\n{fname}")

            # Опционально: открыть файл
            if QMessageBox.question(
                self, "Открыть файл?",
                "Открыть созданный документ Word?",
                QMessageBox.Yes | QMessageBox.No
            ) == QMessageBox.Yes:
                import os
                os.startfile(fname)  # Windows

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить:\n{str(e)}")


    def _prepare_html_for_pandoc(self, html_content):
        """
        Более агрессивный подход: заменяем стили на старые HTML-атрибуты
        """
        import re
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Обрабатываем все ячейки с цветным фоном
        for cell in soup.find_all(['td', 'th']):
            # Проверяем inline style
            if cell.has_attr('style'):
                style = cell['style']
                
                # Извлекаем hex-цвет
                match = re.search(r'background(?:-color)?:\s*#([0-9a-fA-F]{6})', style)
                if match:
                    hex_color = match.group(1)
                    
                    # Конвертируем в RGB для Word
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    
                    # Используем bgcolor (старый HTML атрибут, который Pandoc понимает лучше)
                    cell['bgcolor'] = f"#{hex_color}"
                    
                    # Также оборачиваем содержимое в span с цветом
                    content = cell.decode_contents()
                    cell.clear()
                    new_span = soup.new_tag('span', style=f'background-color:#{hex_color};display:block;padding:2px;')
                    new_span.append(BeautifulSoup(content, 'html.parser'))
                    cell.append(new_span)
            
            # Обрабатываем классы
            if cell.has_attr('class'):
                classes = cell['class']
                
                if 'diag' in classes:
                    cell['bgcolor'] = '#e8e8e8'
                elif 'diag-header' in classes:
                    cell['bgcolor'] = '#b3e0ff'
                elif 'row-header' in classes:
                    cell['bgcolor'] = '#f0f4ff'
        
        return str(soup)


    def _generate_extended_report(self):
        """Генерация расширенного HTML-отчёта — версия с крупным названием и полусерым диагональным текстом"""
        
        # Шаг 1: Получаем список индексов признаков из self.stat_corr
        num_features = len(self.stat_corr.column_names)
        if num_features < 2:
            return "<html><body><h2>Ошибка: выберите хотя бы 2 валидных признака</h2></body></html>"

        selected_indices = list(range(num_features))  # Индексы 0..num_features-1, соответствующие self.stat_corr.column_names

        BLOCK_SIZE = 8

        # Шаг 2: HTML-заголовок и стили (обновлены размеры и цвета для ch10-стиля)
        lines = [
            "<!DOCTYPE html>",
            "<html lang='ru'>",
            "<head>",
            "  <meta charset='UTF-8'>",
            "  <title>Отчет программы MapCor</title>",
            "  <style>",
            "    body {font-family: 'Segoe UI', Arial, sans-serif; margin: 32px; background: #f9f9f9; color: #222; font-size: 1.05em;}",
            "    h1 {color: #1a3c5e; text-align: center; font-size: 2.3em; margin-bottom: 0.5em;}",
            "    h2 {color: #2c5282; text-align: center; font-size: 1.7em; margin: 2em 0 0.8em;}",
            "    .feature-caption {text-align: center; margin: 1.8em 0 0.9em;}",
            "    .feature-caption .name {font-size: 2.4em; font-weight: bold; color: #0f2a6e; letter-spacing: -0.5px;}",
            "    .feature-caption .stats {font-size: 1.05em; font-weight: bold; color: #555; margin-left: 28px;}",
            "    table {border-collapse: collapse; margin: 0 auto 2.4em auto; width: auto; box-shadow: 0 2px 8px rgba(0,0,0,0.07);}",
            "    th, td {border: 1px solid #d0d0d0; padding: 11px 15px; text-align: center; font-size: 1.1em;;-webkit-print-color-adjust: exact; color-adjust: exact;}",
            "    th {background: #e8f0ff; color: #1e3a8a; font-weight: 600;}",
            "    td {font-weight: bold;}",
            "    .diag-header {",
            "      font-size: 1.25em !important;",
            "      font-weight: bold !important;",
            "      background: #b3e0ff !important;",
            "      border: 2px solid #80c0ff !important;",
            "      padding: 12px 16px !important;",
            "    }",
            "    .diag {",
            "      background: #e8e8e8 !important;",
            "      color: #B3B3B3 !important;",           # ← полусерый текст на диагонали
            "      font-style: italic;",
            "      font-weight: bold;",
            "      font-size: 1.0em !important;",        # ← тот же размер, что и остальные значения
            "    }",
            "    .na {color: #888; font-style: italic; font-weight: normal;}",
            "    .row-header {background: #f0f4ff; font-weight: bold; text-align: left; min-width: 60px;}",
            "    td.num {font-family: Consolas, 'Courier New', monospace;-webkit-print-color-adjust: exact; color-adjust: exact;}",
            "    hr {border: 0; height: 1px; background: #ddd; margin: 2.4em 0;}",
            "  </style>",
            "</head>",
            "<body>",
            "<h1>Корреляционный анализ</h1>",
            f"<p align='center' style='font-size:1.25em; margin-bottom:2.2em;'>",
            f"<b>Программа:</b> MapCor ;  ",
            f"<b>Файл:</b> {Path(self.data.filename).name} ;  ",
            f"<b>Число объектов:</b> {self.data.get_count_record()} ;  ",
            f"<b>Число характеристик:</b> {num_features} ;  ",
            
            
        ]

        # Шаг 3: Общая статистика — первая
        lines.append('    <h2>Общая статистика по всем выбранным парам</h2>')
        lines.append('    <table style="width:72%; max-width:950px;">')
        lines.append('      <tr><th>Показатель</th><th>Минимум</th><th>Максимум</th><th>Среднее</th></tr>')
        corr_stat = self.stat_corr.all_pairs_stat['corr']
        lines.append(f'      <tr><td><b>R</b></td><td>{corr_stat.min:.3f}</td><td>{corr_stat.max:.3f}</td><td>{corr_stat.mean:.3f}</td></tr>')
        #dist10_stat = self.stat_corr.all_pairs_stat['dist10']
        #lines.append(f'      <tr><td><b>DIST_10</b></td><td>{dist10_stat.min:.1f}</td><td>{dist10_stat.max:.1f}</td><td>{dist10_stat.mean:.1f}</td></tr>')
        rr_stat = self.stat_corr.all_pairs_stat['rr']
        lines.append(f'      <tr><td><b>RR</b></td><td>{rr_stat.min:.3f}</td><td>{rr_stat.max:.3f}</td><td>{rr_stat.mean:.3f}</td></tr>')
        lines.append('    </table>')
        lines.append('<p style="text-align:center; color:#555; font-size:1.05em; margin: 0.8em 0 2em 0;">')
        lines.append('<b>R</b> — коэффициент ранговой корреляции<br> <b>RR</b> — корреляция корреляций')
        lines.append('</p>')

        #lines.append('    <hr>')

        # Шаг 4: Таблицы по каждой характеристике
        for feature_idx in selected_indices:
            feature_name = self.stat_corr.get_column_name(feature_idx)
            fs = self.stat_corr.feature_stats[feature_idx]

            # Одна строка над таблицей: крупное имя + маленькая статистика
            lines.append(
                '    <div class="feature-caption">'
                f'<span class="name">{feature_name}</span>  '
                f'<span class="stats">'
                f'M(R) = {fs.avg_corr:.3f} ;  '
                #f'M(DIST10) = {fs.avg_dist10:.1f} ;  '
                f'M(RR) = {fs.avg_rr:.3f}'
                f'</span>'
                '</div>'
            )

            # Блочные таблицы
            block_start = 0
            while block_start < len(selected_indices):
                block_end = min(block_start + BLOCK_SIZE - 1, len(selected_indices) - 1)

                lines.append('    <table>')
                lines.append('      <tr><th class="row-header"></th>')

                # Заголовки столбцов — диагональ выделена сильнее
                for j in range(block_start, block_end + 1):
                    col_name = self.stat_corr.get_column_name(selected_indices[j])
                    if selected_indices[j] == feature_idx:
                        lines.append(f'        <th class="diag-header">{col_name}</th>')
                    else:
                        lines.append(f'        <th>{col_name}</th>')
                lines.append('      </tr>')

                # R
                lines.append('      <tr><td class="row-header"><b>R</b></td>')
                for j in range(block_start, block_end + 1):
                    other_idx = selected_indices[j]
                    if feature_idx == other_idx:
                        lines.append('        <td class="diag">1.000</td>')
                    else:
                        pair_idx = self.stat_corr.get_pair_index(feature_idx, other_idx)
                        if pair_idx >= 0:
                            val = self.stat_corr.get_corr(pair_idx)
                            color = get_color_for_r(val)
                            lines.append(f'        <td style="background:{color};" class="num">{val:.3f}</td>')
                        else:
                            lines.append('        <td class="na">—</td>')
                lines.append('      </tr>')

                # DIST_10
               # lines.append('      <tr><td class="row-header"><b>DIST_10</b></td>')
               # for j in range(block_start, block_end + 1):
               #     other_idx = selected_indices[j]
               #     if feature_idx == other_idx:
               #         lines.append('        <td class="diag">100</td>')
               #     else:
               #         pair_idx = self.stat_corr.get_pair_index(feature_idx, other_idx)
               #         if pair_idx >= 0:
               #             val = self.stat_corr.get_dist10(pair_idx)
               #             color = get_color_for_dist10(val)
               #             lines.append(f'        <td style="background:{color};" class="num">{val:.1f}</td>')
               #         else:
               #             lines.append('        <td class="na">—</td>')
               # lines.append('      </tr>')

                # RR
                lines.append('      <tr><td class="row-header"><b>RR</b></td>')
                for j in range(block_start, block_end + 1):
                    other_idx = selected_indices[j]
                    if feature_idx == other_idx:
                        lines.append('        <td class="diag">—</td>')
                    else:
                        pair_idx = self.stat_corr.get_pair_index(feature_idx, other_idx)
                        if pair_idx >= 0:
                            val = self.stat_corr.get_rr(pair_idx)
                            color = get_color_for_rr(val)
                            lines.append(f'        <td style="background:{color};" class="num">{val:.3f}</td>')
                        else:
                            lines.append('        <td class="na">—</td>')
                lines.append('      </tr>')

                lines.append('    </table>')
                block_start = block_end + 1

            lines.append('    <hr>')

        lines.append('  </body></html>')
        return '\n'.join(lines)

    def _generate_old_report(self):
        """
        Генерация классической версии HTML-отчёта с матрицей:
        - выше диагонали → Spearman R
        - ниже диагонали → DIST₁₀
        Работает на основе содержимого self.stat_corr, без зависимости от текущего состояния чекбоксов.
        """
        # Если ничего не посчитано — выходим рано
        if self.stat_corr.count() == 0 or len(self.stat_corr.column_names) < 2:
            return "<html><body><h2>Ошибка: нет рассчитанных корреляций или менее 2 признаков</h2></body></html>"

        # Берём все признаки, которые участвовали в расчёте (из stat_corr)
        selected_indices = list(range(len(self.stat_corr.column_names)))

        BLOCK_SIZE = 8  # можно оставить или убрать блочность — решайте по красоте

        lines = [
            "<!DOCTYPE html>",
            "<html lang='ru'>",
            "<head>",
            "  <meta charset='UTF-8'>",
            "  <title>Отчет программы MapCor — классическая матрица</title>",
            "  <style>",
            "    body {font-family: 'Segoe UI', Arial, sans-serif; margin: 32px; background: #f9f9f9; color: #222; font-size: 1.05em;}",
            "    h1 {color: #1a3c5e; text-align: center; font-size: 2.3em; margin-bottom: 0.5em;}",
            "    h2 {color: #2c5282; text-align: center; font-size: 1.7em; margin: 2em 0 0.8em;}",
            "    table {border-collapse: collapse; margin: 0 auto 3em auto; width: auto; box-shadow: 0 3px 12px rgba(0,0,0,0.08);}",
            "    th, td {border: 1px solid #d0d0d0; padding: 10px 14px; text-align: center; font-size: 1.05em; -webkit-print-color-adjust: exact; color-adjust: exact;}",
            "    th {background: #e8f0ff; color: #1e3a8a; font-weight: 600;}",
            "    td {font-weight: bold;}",
            "    .diag {",
            "      background: #e0e0ff !important;",
            "      color: #888 !important;",
            "      font-style: italic;",
            "      font-weight: bold;",
            "    }",
            "    .na {color: #999; font-style: italic; font-weight: normal;}",
            "    .row-header {background: #f0f4ff; font-weight: bold; text-align: left; min-width: 180px; padding-left: 12px;}",
            "    td.num {font-family: Consolas, 'Courier New', monospace;}",
            "    hr {border: 0; height: 1px; background: #ddd; margin: 3em 0;}",
            "  </style>",
            "</head>",
            "<body>",
            "<h1>Матрица корреляций Спирмена и DIST₁₀</h1>",
            f"<p style='text-align:center; font-size:1.25em; margin-bottom:2.2em;'>",
            f"<b>Файл:</b> {Path(self.data.filename).name} ;  ",
            f"<b>Записей:</b> {self.data.get_count_record()} ;  ",
            f"<b>Признаков в расчёте:</b> {len(selected_indices)} ;  ",
            f"<b>Пар:</b> {self.stat_corr.count()}",
            "</p>",
            "<hr>",
        ]

        # Общая статистика
        lines.append('    <h2>Общая статистика по всем парам</h2>')
        lines.append('    <table style="width:72%; max-width:950px; margin-bottom:2.5em;">')
        lines.append('      <tr><th>Показатель</th><th>Минимум</th><th>Максимум</th><th>Среднее</th></tr>')

        corr_stat = self.stat_corr.all_pairs_stat['corr']
        lines.append(f'      <tr><td><b>R (Spearman)</b></td><td>{corr_stat.min:.3f}</td><td>{corr_stat.max:.3f}</td><td>{corr_stat.mean:.3f}</td></tr>')

        dist10_stat = self.stat_corr.all_pairs_stat['dist10']
        lines.append(f'      <tr><td><b>DIST₁₀</b></td><td>{dist10_stat.min:.1f}</td><td>{dist10_stat.max:.1f}</td><td>{dist10_stat.mean:.1f}</td></tr>')

        rr_stat = self.stat_corr.all_pairs_stat['rr']
        lines.append(f'      <tr><td><b>RR (мета-корр)</b></td><td>{rr_stat.min:.3f}</td><td>{rr_stat.max:.3f}</td><td>{rr_stat.mean:.3f}</td></tr>')

        lines.append('    </table>')
        lines.append('<p style="text-align:center; color:#555; font-size:1.05em; margin: 0.8em 0 2.5em 0;">')
        lines.append('Выше диагонали — <b>R (Spearman)</b><br>Ниже диагонали — <b>DIST₁₀</b>')
        lines.append('</p>')

        # Основная матрица
        lines.append('<table style="margin: 0 auto 4em auto;">')

        # Заголовочная строка
        lines.append('  <tr>')
        lines.append('    <th class="row-header"></th>')
        for col_idx in selected_indices:
            col_name = self.stat_corr.get_column_name(col_idx)
            lines.append(f'    <th title="{col_name}">{col_name}</th>')
        lines.append('  </tr>')

        # Строки матрицы
        for row_i, row_idx in enumerate(selected_indices):
            row_name = self.stat_corr.get_column_name(row_idx)
            lines.append('  <tr>')
            lines.append(f'    <td class="row-header">{row_name}</td>')

            for col_j, col_idx in enumerate(selected_indices):
                if row_i == col_j:
                    # Диагональ — всегда 1.000 для R
                    lines.append('    <td class="diag">1.000</td>')
                elif row_i < col_j:
                    # Выше диагонали → Spearman R
                    pair_idx = self.stat_corr.get_pair_index(row_idx, col_idx)
                    if pair_idx >= 0:
                        val = self.stat_corr.get_corr(pair_idx)
                        color = get_color_for_r(val)
                        lines.append(f'    <td style="background:{color};" class="num">{val:.3f}</td>')
                    else:
                        lines.append('    <td class="na">—</td>')
                else:
                    # Ниже диагонали → DIST10
                    pair_idx = self.stat_corr.get_pair_index(row_idx, col_idx)
                    if pair_idx >= 0:
                        val = self.stat_corr.get_dist10(pair_idx)
                        color = get_color_for_dist10(val)
                        lines.append(f'    <td style="background:{color};" class="num">{val:.1f}</td>')
                    else:
                        lines.append('    <td class="na">—</td>')

            lines.append('  </tr>')

        lines.append('</table>')
        lines.append('<hr style="margin: 4em 0 2em 0;">')

        lines.append('  </body></html>')
        return '\n'.join(lines)

    def _generate_stats_report(self, selected_columns=None):
        """
        Генерирует HTML-отчёт по статистике признаков.
        Подсветка заголовков убрана, добавлена возможность включать/выключать столбцы и пункты легенды.
        """
        if self.data.df is None or self.data.df.empty:
            return "<h2 style='text-align:center;color:#c53030;'>Нет загруженных данных</h2>"

        stats_df = self.data.get_full_statistics()
        if stats_df is None or stats_df.empty:
            return "<h2 style='text-align:center;color:#c53030;'>Нет числовых признаков</h2>"

        # ────────────────────────────────────────────────────────────────
        # Настройки: какие столбцы показывать в таблице
        # ────────────────────────────────────────────────────────────────
        SHOW_COLUMNS = {
            'count'                  : False,
            'nan_percent'            : False,
            'min'                    : True,
            'repeating_min_percent'  : True,
            'below_lod_percent'      : False,
            'zero_percent'           : False,
            '5%'                     : False,
            'Q1'                     : False,
            'median'                 : True,
            'Q3'                     : False,
            '95%'                    : False,
            'max'                    : True,
            'mean'                   : True,
            'std'                    : True,
            'CV_percent'             : True,
            'variance'               : False,
            'skew'                   : False,   # скрыт по умолчанию
            'kurtosis'               : False,  # скрыт по умолчанию
            'unique_count'           : False,
            'J'                      : True,
        }

        # ────────────────────────────────────────────────────────────────
        # Настройки: какие пункты показывать в легенде
        # ────────────────────────────────────────────────────────────────
        SHOW_LEGEND_ITEMS = {
            'count'                  : False,
            'nan_percent'            : False,
            'min_max'                : True,
            'repeating_min_percent'  : True,
            'below_lod_percent'      : False,
            'zero_percent'           : False,
            'percentiles'            : False,
            'quartiles_median'       : False,
            'mean_std'               : True,
            'CV_percent'             : True,
            'variance'               : False,
            'skew_kurtosis'          : False,
            'unique_count'           : False,
            'J'                      : True,
        }

        # Фильтруем столбцы, которые хотим показать
        columns_to_show = [col for col in stats_df.columns if SHOW_COLUMNS.get(col, False)]
        if not columns_to_show:
            return "<h2 style='text-align:center;color:#c53030;'>Нет выбранных для отображения статистик</h2>"

        stats_df = stats_df[columns_to_show].copy()
        stats_df.index.name = 'Признак'

        # Дополнительная фильтрация по выбранным признакам (если передан список индексов)
        if selected_columns:
            valid_names = [self.data.get_column_name(i) for i in selected_columns
                           if 0 <= i < len(self.data.df.columns)]
            stats_df = stats_df.loc[stats_df.index.isin(valid_names)]

        if stats_df.empty:
            return "<h2 style='text-align:center;color:#c53030;'>Нет выбранных числовых признаков</h2>"

        # Форматирование значений для отображения
        display_df = stats_df.copy()
        for col in display_df.columns:
            if col in ['min', '5%', 'Q1', 'median', 'Q3', '95%', 'max', 'mean', 'std']:
                display_df[col] = display_df[col].map(lambda x: f"{x:.3f}" if pd.notna(x) else "—")
            elif col in ['CV_percent', 'below_lod_percent', 'repeating_min_percent', 'zero_percent', 'nan_percent']:
                display_df[col] = display_df[col].map(lambda x: f"{x:.1f}" if pd.notna(x) else "—")
            elif col == 'variance':
                display_df[col] = display_df[col].map(lambda x: f"{x:.6f}" if pd.notna(x) else "—")
            elif col == 'J':
                display_df[col] = display_df[col].map(lambda x: f"{x:.3f}" if pd.notna(x) else "—")
            else:
                display_df[col] = display_df[col].astype(str).replace('nan', '—')

        # ────────────────────────────────────────────────────────────────
        # HTML-отчёт
        # ────────────────────────────────────────────────────────────────
        lines = [
            "<!DOCTYPE html>",
            "<html lang='ru'>",
            "<head>",
            "<meta charset='UTF-8'>",
            "<title>Статистический отчёт — MAPCOR-P</title>",
            "<style>",
            "  body {font-family: 'Segoe UI', Arial, sans-serif; margin:0; padding:20px; background:#f8fafc; color:#1e293b; line-height:1.6;}",
            "  .container {max-width:1480px; margin:0 auto; background:white; padding:30px; border-radius:12px; box-shadow:0 10px 30px rgba(0,0,0,0.08);}",
            "  h1 {text-align:center; color:#1e40af; margin-bottom:8px;}",
            "  .subtitle {text-align:center; color:#475569; font-size:1.1em; margin-bottom:30px;}",
            "  table {width:100%; border-collapse:collapse; margin:25px 0; font-size:0.94em;}",
            "  th {background:#f1f5f9; color:#334155; padding:9px 8px; text-align:center; font-weight:600; border:1px solid #e2e8f0;}",
            "  td {padding:8px 10px; border:1px solid #e2e8f0; text-align:right;}",
            "  .row-header {text-align:left !important; font-weight:600; background:#f8fafc; min-width:60px; padding-left:6px;}",
            "  .lod-col   {background:#fefce8;}",
            "  .cv-col    {background:#fff7ed;}",
            "  .j-col     {background:#f0fdf4; font-weight:bold;}",
            "  .percentile{background:#f8fafc;}",
            "  .na {color:#94a3b8; font-style:italic;}",
            #"  .table-wrapper {overflow-x:auto; margin:30px 0; padding:10px; background:#f8fafc; border-radius:8px;}",
            "  hr {border:none; height:1px; background:#e2e8f0; margin:40px 0;}",
            "</style>",
            "</head>",
            "<body>",
            "<div class='container'>",
            f"<h1>Статистический отчёт по признакам</h1>",
            f"<p class='subtitle'>Файл: <b>{Path(self.data.filename).name}</b> | "
            f"Записей: <b>{self.data.get_count_record():,}</b> | "
            f"Признаков: <b>{len(stats_df)}</b></p>",
            "<hr>",
        ]

        ROWS_PER_TABLE = 250
        chunks = [display_df.iloc[i:i+ROWS_PER_TABLE] for i in range(0, len(display_df), ROWS_PER_TABLE)]

        for idx, chunk in enumerate(chunks, 1):
           # lines.append("<div class='table-wrapper'>")
            lines.append("<table>")
            
            # Заголовки
            lines.append("<tr>")
            lines.append("<th class='row-header'>Признак</th>")
            for col in chunk.columns:
                title_map = {
                    'repeating_min_percent': 'Мин. повт., %',
                    'below_lod_percent'    : '≤LOD, %',
                    'zero_percent'         : 'Нули, %',
                    'CV_percent'           : 'CV, %',
                    'nan_percent'          : 'NaN, %',
                    'unique_count'         : 'Уник.',
                    'variance'             : 'Var',
                    'J'                    : 'J (информ.)'
                }
                display_name = title_map.get(col, col)
                cls = ""
                if col in ['5%', 'Q1', 'Q3', '95%']: cls = " class='percentile'"
                if col == 'below_lod_percent':       cls = " class='lod-col'"
                if col == 'CV_percent':              cls = " class='cv-col'"
                if col == 'J':                       cls = " class='j-col'"
                lines.append(f"<th{cls}>{display_name}</th>")
            lines.append("</tr>")

            # Данные
            for feature, row in chunk.iterrows():
                lines.append("<tr>")
                lines.append(f"<td class='row-header'>{feature}</td>")
                for val_str in row:
                    lines.append(f"<td>{val_str}</td>")
                lines.append("</tr>")
            
            lines.append("</table>")
            if len(chunks) > 1:
                lines.append(f"<p style='text-align:right; color:#64748b; font-size:0.9em;'>Таблица {idx} из {len(chunks)}</p>")
            #lines.append("</div>")

        # ────────────────────────────────────────────────────────────────
        # Легенда — только включённые пункты
        # ────────────────────────────────────────────────────────────────
        legend_lines = []
        if SHOW_LEGEND_ITEMS.get('count'):
            legend_lines.append("  <li><strong>count</strong> — количество непропущенных значений</li>")
        if SHOW_LEGEND_ITEMS.get('nan_percent'):
            legend_lines.append("  <li><strong>NaN, %</strong> — доля пропущенных значений</li>")
        if SHOW_LEGEND_ITEMS.get('min_max'):
            legend_lines.append("  <li><strong>min / max</strong> — минимальное и максимальное значение</li>")
        if SHOW_LEGEND_ITEMS.get('repeating_min_percent'):
            legend_lines.append("  <li><strong>Мин. повт., %</strong> — сколько процентов строк имеют значение, равное минимальному</li>")
        if SHOW_LEGEND_ITEMS.get('below_lod_percent'):
            legend_lines.append("  <li><strong>≤LOD, %</strong> — доля значений ≤ 0.03 (включая NaN)</li>")
        if SHOW_LEGEND_ITEMS.get('zero_percent'):
            legend_lines.append("  <li><strong>Нули, %</strong> — доля нулевых или почти нулевых значений (≤ 0.03)</li>")
        if SHOW_LEGEND_ITEMS.get('percentiles'):
            legend_lines.append("  <li><strong>5% / 95%</strong> — 5-й и 95-й перцентили</li>")
        if SHOW_LEGEND_ITEMS.get('quartiles_median'):
            legend_lines.append("  <li><strong>Q1 / median / Q3</strong> — квартили и медиана</li>")
        if SHOW_LEGEND_ITEMS.get('mean_std'):
            legend_lines.append("  <li><strong>mean / std</strong> — среднее и стандартное отклонение</li>")
        if SHOW_LEGEND_ITEMS.get('CV_percent'):
            legend_lines.append("  <li><strong>CV, %</strong> — коэффициент вариации = (std / |mean|) × 100 %</li>")
        if SHOW_LEGEND_ITEMS.get('variance'):
            legend_lines.append("  <li><strong>Var</strong> — дисперсия</li>")
        if SHOW_LEGEND_ITEMS.get('skew_kurtosis'):
            legend_lines.append("  <li><strong>skew / kurtosis</strong> — асимметрия и эксцесс</li>")
        if SHOW_LEGEND_ITEMS.get('unique_count'):
            legend_lines.append("  <li><strong>Уник.</strong> — количество уникальных значений</li>")
        if SHOW_LEGEND_ITEMS.get('J'):
            legend_lines.append("  <li><strong>J (информ.)</strong> — нормированная информативность по Шеннону (6 фиксированных интервалов)<br>"
                                "    · <strong>J ≈ 1.0</strong> — почти все значения в одном интервале → монолитный пласт<br>"
                                "    · <strong>J ≈ 0.0</strong> — равномерное распределение по всем 6 интервалам → максимальная гетерогенность<br>"
                                "    · Рекомендуемый порог однородности: <strong>J ≥ 0.65</strong></li>")

        if legend_lines:
            lines.extend([
                "<hr>",
                "<div style='background:#f8fafc; padding:24px; border-radius:10px; font-size:0.98em; line-height:1.7;'>",
                "<h3 style='color:#1e40af; margin:0 0 16px 0;'>Расшифровка статистических показателей</h3>",
                "<ul style='margin:0; padding-left:20px; columns:2; column-gap:40px;'>",
            ])
            lines.extend(legend_lines)
            lines.extend([
                "</ul>",
                "</div>",
            ])

        lines.append("</div></body></html>")
        return "\n".join(lines)

    def act_save_result(self):
        if self.stat_corr.count() == 0:
            QMessageBox.information(self, "Нет данных", "Нет рассчитанных результатов для сохранения.")
            return

        # Диалог сохранения — по умолчанию .txt
        fname, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить таблицу результатов",
            os.path.join(self._get_initial_dir(), "results.txt"),
            "Текстовые файлы (*.txt);;Все файлы (*.*)"
        )
        if not fname:
            return

        # Обновляем рабочую директорию на директорию сохраненного файла
        self.working_directory = os.path.dirname(fname)

        # Добавляем расширение .txt, если пользователь его не указал
        if not fname.lower().endswith('.txt'):
            fname += '.txt'

        try:
            with open(fname, "w", encoding="utf-8") as f:
                        # 1. Первая строка — количество пар (а не количество признаков!)
                        f.write(f"4\n")

                        # 2. Заголовок таблицы
                        f.write("n\tname\tR\tRR\n")

                        # 3. Данные по всем парам
                        for i in range(self.stat_corr.count()):
                            pair_name = self.stat_corr.get_pair_name(i)
                            r_value   = self.stat_corr.get_corr(i)
                            rr_value  = self.stat_corr.get_rr(i)

                            # Порядковый номер начиная с 1
                            num = i + 1

                            # Форматирование значений с учётом возможных NaN
                            r_str  = f"{r_value:.3f}"  if not np.isnan(r_value)  else "—"
                            rr_str = f"{rr_value:.3f}" if not np.isnan(rr_value) else "—"

                            line = f"{num}\t{pair_name}\t{r_str}\t{rr_str}"
                            f.write(line + "\n")

            QMessageBox.information(self, "Сохранено", f"Результаты сохранены в файл:\n{fname}")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка сохранения",
                                f"Не удалось сохранить файл:\n{str(e)}")
            

if __name__ == "__main__":
    app = QApplication(sys.argv)

# Самый современный и чистый вид на Windows 10/11
    app.setStyle("Fusion")

    # Лёгкая тёмная/светлая тема с хорошей читаемостью
    app.setStyleSheet("""
        QMainWindow {
            background-color: #f8f9fa;
        }
        QGroupBox {
            font-weight: bold;
            border: 1px solid #ced4da;
            border-radius: 4px;
            margin-top: 10px;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px;
        }
        QTableView {
            gridline-color: #dee2e6;
            alternate-background-color: #f1f3f5;
        }
        QHeaderView::section {
            background-color: #e9ecef;
            padding: 6px;
            border: 1px solid #ced4da;
        }
        QListWidget {
            border: 1px solid #ced4da;
            border-radius: 4px;
            background-color: white;
        }
        QStatusBar {
            background-color: #e9ecef;
            color: #495057;
        }
    """)

    # Шрифт (очень важно для профессионального вида)
    font = QFont("Segoe UI", 12)
    app.setFont(font)

    window = MainWindow()
    window.show()
    sys.exit(app.exec())